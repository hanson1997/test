#!/usr/bin/env python3
"""Generate the Technical Department BIA register as a Word document.

Follows the original Activities & Potential Risk Register layout:
title, short intro, one landscape table, header row repeating on each page.

Kept: Activity, Potential Risk
Removed: Seasonal/Peak, Interdependent Departments, 1-Hour columns,
         Impact Description @ 1 Day
Skipped: Financial Impact per Day ($)
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT

from bia_register_data import CURRENT_STATE, DEPENDENCIES, ROWS, TEST_RESULTS

NAVY = "1B365D"
TEAL = "1F7A8C"
WHITE = "FFFFFF"
BODY = "1A1A1A"
RATING_FILL = {
    "5 - Critical": "C0392B",
    "4 - Major": "F4D03F",
    "3 - Moderate": "E67E22",
    "2 - Minor": "7FB3D5",
    "1 - Insignificant": "82C785",
}
RATING_FONT = {
    "5 - Critical": WHITE,
    "4 - Major": BODY,
    "3 - Moderate": BODY,
    "2 - Minor": BODY,
    "1 - Insignificant": BODY,
}
HEADERS = [
    "Activity",
    "Potential Risk",
    "Impact Rating @ 1 Day",
    "Impact Description @ 1 Week",
    "Impact Rating @ 1 Week",
    "Client / Customer Impact",
    "Overall Criticality Ranking",
    "MTPD (Hrs)",
    "RTO (Hrs)",
    "RPO",
]
# Share of usable page width (A3 landscape, ~40.1 cm after 1.2 cm margins)
COL_WIDTHS_CM = [3.4, 4.6, 3.0, 5.4, 3.0, 4.8, 3.2, 2.2, 2.2, 2.3]
RATING_COLS = {2, 4, 6}
CENTER_COLS = {2, 4, 6, 7, 8, 9}

DEP_HEADERS = [
    "Activity",
    "Potential Risk",
    "Upstream Dependencies (Relies On)",
    "Downstream Dependencies (Relied On By)",
    "Key Systems / Applications",
    "Key Data / Records",
    "Key Personnel (Primary + Backup)",
    "Key Vendors / Third Parties",
]
STRATEGY_HEADERS = [
    "Activity",
    "Potential Risk",
    "Single Points of Failure",
    "Existing Safeguards / Controls",
    "Recommended Recovery Strategy",
    "Alternate Site/Access Required?",
    "Cross-Trained Backup Staff Identified?",
]
TEST_HEADERS = [
    "Activity",
    "Potential Risk",
    "Test Result Summary",
]
TEST_WIDTHS_CM = [6.0, 12.0, 16.0]


def set_run(run, *, size=10, bold=False, color=BODY, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="8AA0B4", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def valign_center(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:vAlign"):
            tcPr.remove(child)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tcPr.append(v)


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcMar"):
            tcPr.remove(child)
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    trPr.append(hdr)


def set_table_fixed(table, widths):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn("w:tblW"):
            tblPr.remove(child)
        if child.tag == qn("w:tblLayout"):
            tblPr.remove(child)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    total = int(sum(widths) * 567)  # cm -> twips (approx 567 twips/cm)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def fill_cell(cell, text, *, bold=False, size=8, fill=None, font_color=BODY, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=font_color)
    valign_center(cell)
    set_cell_margins(cell)
    set_cell_borders(cell)
    if fill:
        shade_cell(cell, fill)


def validate_logic():
    scale_rank = {
        "1 - Insignificant": 1,
        "2 - Minor": 2,
        "3 - Moderate": 3,
        "4 - Major": 4,
        "5 - Critical": 5,
    }
    problems = []
    for row in ROWS:
        if scale_rank[row["r1w"]] < scale_rank[row["r1d"]]:
            problems.append(f"{row['activity']}: 1-week rating lower than 1-day")
        if row["overall"] != row["r1w"]:
            problems.append(f"{row['activity']}: overall does not match 1-week")
        if not (row["mtpd"] > row["rto"]):
            problems.append(f"{row['activity']}: MTPD is not greater than RTO")
    missing = [row["activity"] for row in ROWS if row["activity"] not in DEPENDENCIES]
    extra = sorted(set(DEPENDENCIES) - {row["activity"] for row in ROWS})
    if missing:
        problems.append("Missing dependencies for: " + "; ".join(missing))
    if extra:
        problems.append("Unexpected dependency keys: " + "; ".join(extra))
    missing_cs = [row["activity"] for row in ROWS if row["activity"] not in CURRENT_STATE]
    extra_cs = sorted(set(CURRENT_STATE) - {row["activity"] for row in ROWS})
    if missing_cs:
        problems.append("Missing current-state rows for: " + "; ".join(missing_cs))
    missing_tr = [row["activity"] for row in ROWS if row["activity"] not in TEST_RESULTS]
    extra_tr = sorted(set(TEST_RESULTS) - {row["activity"] for row in ROWS})
    if extra_cs:
        problems.append("Unexpected current-state keys: " + "; ".join(extra_cs))
    if missing_tr:
        problems.append("Missing test-result rows for: " + "; ".join(missing_tr))
    if extra_tr:
        problems.append("Unexpected test-result keys: " + "; ".join(extra_tr))
    if problems:
        raise SystemExit("Logic errors:\n- " + "\n- ".join(problems))


def add_intro(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Technical Department – Activities & Potential Risk Register")
    set_run(run, size=16, bold=True, color=NAVY)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    intro.paragraph_format.space_before = Pt(0)
    run = intro.add_run(
        "This register maps each functional activity area performed by the Technical "
        "Department to its potential risk, impact over time (1 day and 1 week), "
        "client/customer impact, overall criticality, and recovery requirements. "
        "Financial Impact per Day ($) is omitted. Ratings use the existing 1 – "
        "Insignificant to 5 – Critical scale. Figures are a workshop draft reasoned "
        "from the previous 1-hour / 1-day register — confirm MTPD, RTO and RPO "
        "against contracts and tested restore times before treating them as final."
    )
    set_run(run, size=10, color=BODY)

    key = doc.add_paragraph()
    key.paragraph_format.space_after = Pt(10)
    items = [
        ("MTPD", "Maximum Tolerable Period of Disruption (hours) — the point beyond which loss of the function threatens the business itself. Always the largest of the three (MTPD > RTO > RPO in urgency)."),
        ("RTO", "Recovery Time Objective (hours) — the maximum time allowed to restore the function to minimum service."),
        ("RPO", "Recovery Point Objective — the maximum tolerable data loss, expressed as time (e.g. “15 minutes”, “1 business day”). N/A where the function is process-only."),
    ]
    for i, (abbr, meaning) in enumerate(items):
        if i:
            key.add_run("  ")
        r = key.add_run(f"{abbr}: ")
        set_run(r, size=9, bold=True, color=NAVY)
        r = key.add_run(meaning)
        set_run(r, size=9, color="334155")
        if i < len(items) - 1:
            key.add_run("  |  ")


def merge_fill(cells, text, fill):
    cells[0].merge(cells[-1])
    fill_cell(cells[0], text, bold=True, size=9, fill=fill, font_color=WHITE, center=True)


def build_table(doc):
    table = doc.add_table(rows=2 + len(ROWS), cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table, COL_WIDTHS_CM)

    group = table.rows[0]
    repeat_header(group)
    group.height = Cm(0.7)
    merge_fill(group.cells[0:2], "FUNCTION", NAVY)
    merge_fill(group.cells[2:7], "IMPACT ANALYSIS OVER TIME", TEAL)
    merge_fill(group.cells[7:10], "RECOVERY REQUIREMENTS", NAVY)
    for i in range(len(HEADERS)):
        group.cells[i].width = Cm(COL_WIDTHS_CM[i])

    hdr = table.rows[1]
    repeat_header(hdr)
    hdr.height = Cm(1.05)
    for i, label in enumerate(HEADERS):
        fill_cell(
            hdr.cells[i],
            label,
            bold=True,
            size=8,
            fill=TEAL if 2 <= i <= 6 else NAVY,
            font_color=WHITE,
            center=True,
        )
        hdr.cells[i].width = Cm(COL_WIDTHS_CM[i])

    for r_i, row in enumerate(ROWS):
        cells = table.rows[r_i + 2].cells
        table.rows[r_i + 2].height = Cm(2.05)
        values = [
            row["activity"],
            row["risk"],
            row["r1d"],
            row["d1w"],
            row["r1w"],
            row["client"],
            row["overall"],
            row["mtpd"],
            row["rto"],
            row["rpo"],
        ]
        zebra = "F4F7FA" if r_i % 2 else WHITE
        for c_i, value in enumerate(values):
            if c_i in RATING_COLS:
                rating = str(value)
                fill_cell(
                    cells[c_i],
                    rating,
                    bold=True,
                    size=8,
                    fill=RATING_FILL[rating],
                    font_color=RATING_FONT[rating],
                    center=True,
                )
            else:
                fill_cell(
                    cells[c_i],
                    value,
                    bold=(c_i == 0),
                    size=8,
                    fill=zebra,
                    font_color=BODY,
                    center=c_i in CENTER_COLS,
                )
            cells[c_i].width = Cm(COL_WIDTHS_CM[c_i])

    # Column grid widths on tblGrid
    tbl = table._tbl
    grid = tbl.tblGrid
    for i, gw in enumerate(grid.gridCol_lst):
        gw.set(qn("w:w"), str(int(COL_WIDTHS_CM[i] * 567)))


def add_footer_note(doc):
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    run = note.add_run(
        "Draft for Technical Department workshop. Impact Rating @ 1 Day, Impact Rating "
        "@ 1 Week, and Overall Criticality Ranking use the same 1 – Insignificant → "
        "5 – Critical values. Overall ranking follows the 1-week rating. 1-week ratings "
        "are never lower than 1-day ratings."
    )
    set_run(run, size=9, color="5A6A7A")


def add_dependencies_intro(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Technical Department – Activities & Potential Risk Register")
    set_run(run, size=16, bold=True, color=NAVY)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    run = intro.add_run(
        "This register maps each functional activity area performed by the Technical "
        "Department to its potential risk and to a Test Result Summary for BCP/DR "
        "exercises. Only Activity and Potential Risk are kept from the original register. "
        "No completed Technical Department restore/failover test log was provided, so "
        "these entries do not invent a pass. Rows are either not yet tested (with a "
        "schedule) or they note that remote/VPN work is already BAU while a formal "
        "test still has not been run. Replace with the real result after each exercise."
    )
    set_run(run, size=10, color=BODY)

    key = doc.add_paragraph()
    key.paragraph_format.space_after = Pt(10)
    items = [
        ("Test Result Summary", "Short note on the last BCP/DR exercise for this activity — a completed result with any finding, or “Not yet tested” plus a scheduled window. Same style as the organisation sample."),
    ]
    for i, (abbr, meaning) in enumerate(items):
        r = key.add_run(f"{abbr}: ")
        set_run(r, size=9, bold=True, color=NAVY)
        r = key.add_run(meaning)
        set_run(r, size=9, color="334155")
        if i < len(items) - 1:
            sep = key.add_run("  |  ")
            set_run(sep, size=9, color="94A3B8")


def build_dependencies_table(doc):
    table = doc.add_table(rows=2 + len(ROWS), cols=len(TEST_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table, TEST_WIDTHS_CM)

    group = table.rows[0]
    repeat_header(group)
    group.height = Cm(0.7)
    merge_fill(group.cells[0:2], "FUNCTION", NAVY)
    fill_cell(group.cells[2], "TEST RESULTS", bold=True, size=9, fill=NAVY, font_color=WHITE, center=True)
    for i in range(len(TEST_HEADERS)):
        group.cells[i].width = Cm(TEST_WIDTHS_CM[i])

    hdr = table.rows[1]
    repeat_header(hdr)
    hdr.height = Cm(1.05)
    for i, label in enumerate(TEST_HEADERS):
        fill_cell(
            hdr.cells[i],
            label,
            bold=True,
            size=9,
            fill=NAVY,
            font_color=WHITE,
            center=True,
        )
        hdr.cells[i].width = Cm(TEST_WIDTHS_CM[i])

    for r_i, row in enumerate(ROWS):
        cells = table.rows[r_i + 2].cells
        table.rows[r_i + 2].height = Cm(1.7)
        values = [
            row["activity"],
            row["risk"],
            TEST_RESULTS[row["activity"]],
        ]
        zebra = "F4F7FA" if r_i % 2 else WHITE
        for c_i, value in enumerate(values):
            fill_cell(
                cells[c_i],
                value,
                bold=(c_i == 0),
                size=9,
                fill=zebra,
                font_color=BODY,
                center=False,
            )
            cells[c_i].width = Cm(TEST_WIDTHS_CM[c_i])

    tbl = table._tbl
    grid = tbl.tblGrid
    for i, gw in enumerate(grid.gridCol_lst):
        gw.set(qn("w:w"), str(int(TEST_WIDTHS_CM[i] * 567)))


def add_dependencies_note(doc):
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    run = note.add_run(
        "Draft for workshop. These summaries are not a record of tests that have "
        "already been run. After each exercise, replace the row with the real result "
        "in the sample style (e.g. “Failover successful; manual logging step took "
        "longer than planned” or “Remote-work activation smooth; no issues identified”)."
    )
    set_run(run, size=9, color="5A6A7A")


def main():
    validate_logic()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # A3 landscape so the wide register stays readable, matching the original multi-column table.
    section.page_width = Cm(42.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(10)

    add_dependencies_intro(doc)
    build_dependencies_table(doc)
    add_dependencies_note(doc)

    path = "/workspace/Technical_Department_BIA_Register.docx"
    doc.save(path)
    print(f"Wrote {path} ({len(ROWS)} activities, test-result column only)")


if __name__ == "__main__":
    main()
