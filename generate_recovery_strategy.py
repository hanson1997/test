#!/usr/bin/env python3
"""Fill the Departmental Recovery Strategy Register from the Technical BIA."""

import csv
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.workbook.defined_name import DefinedName
from openpyxl.worksheet.datavalidation import DataValidation
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BIA_PATH = Path(
    "/home/ubuntu/.cursor/projects/workspace/uploads/"
    "Departmental_BIA_Register_Technical_260c.csv"
)
TEMPLATE_PATH = Path(
    "/home/ubuntu/.cursor/projects/workspace/uploads/"
    "2_-_Recovery_Strategy_Register_2_1e42.csv"
)
OUT_CSV = Path("/workspace/Technical_Department_Recovery_Strategy_Register.csv")
OUT_XLSX = Path("/workspace/Technical_Department_Recovery_Strategy_Register.xlsx")
OUT_DOCX = Path("/workspace/Technical_Department_Recovery_Strategy_Register.docx")

NAVY = "1B365D"
TEAL = "1F7A8C"
WHITE = "FFFFFF"
BODY = "1A1A1A"

# Must match the Excel dropdown exactly — no other combinations.
ALLOWED_STRATEGY_TYPES = {
    "People",
    "Process",
    "Technology",
    "Facilities",
    "People + Process",
    "People + Technology",
    "Technology + Facilities",
    "People + Process + Technology",
}

# Chosen from function name, overall impact, and the BIA recovery summary.
# Facilities is not used: Technical recovery is remote/VPN and systems restore,
# not a second building (that dropdown value is for workplace/site strategies).
STRATEGY_TYPE = {
    "Requirements & Solution Design": "People + Technology",
    "Application Development (Coding)": "People + Technology",
    "Source Code & Version Control": "People + Process + Technology",
    "Code Review": "People + Process",
    "Testing & Quality Assurance": "People + Technology",
    "Build Management": "People + Technology",
    "CI/CD Pipeline Management": "People + Process + Technology",
    "Release Deployment (Execution)": "People + Process + Technology",
    "Database & Data Structure Management": "People + Process + Technology",
    "Third-Party & Dependency Management": "People + Technology",
    "Security & Access Configuration": "People + Process + Technology",
    "Environment Management": "People + Technology",
    "Data Handling & ETL": "People + Process + Technology",
    "Production Support & Troubleshooting": "People + Process + Technology",
    "Access & Account Management": "People + Process + Technology",
    "Release Management": "People + Process",
    "Cloud Infrastructure Management": "People + Technology",
    "Logging & Monitoring": "People + Technology",
    "Incident Management": "People + Process + Technology",
    "Decommissioning": "People + Process",
}

# Resource packs already mentioned in the BIA strategy / safeguards.
RESOURCES_DOCUMENTED = {
    "Requirements & Solution Design": "No",
    "Application Development (Coding)": "No",
    "Source Code & Version Control": "No",
    "Code Review": "No",
    "Testing & Quality Assurance": "Yes",
    "Build Management": "Yes",
    "CI/CD Pipeline Management": "Yes",
    "Release Deployment (Execution)": "Yes",
    "Database & Data Structure Management": "Yes",
    "Third-Party & Dependency Management": "Yes",
    "Security & Access Configuration": "Yes",
    "Environment Management": "Yes",
    "Data Handling & ETL": "Yes",
    "Production Support & Troubleshooting": "Yes",
    "Access & Account Management": "Yes",
    "Release Management": "No",
    "Cloud Infrastructure Management": "Yes",
    "Logging & Monitoring": "Yes",
    "Incident Management": "Yes",
    "Decommissioning": "Yes",
}


def clean(value):
    return (
        str(value)
        .replace("\xa0", " ")
        .replace("\n", " ")
        .replace("\r", " ")
        .replace("  ", " ")
        .strip()
    )


def load_bia():
    with BIA_PATH.open(newline="", encoding="latin-1") as f:
        rows = list(csv.reader(f))
    headers = [h.strip() for h in rows[3]]
    idx = {h: i for i, h in enumerate(headers)}
    records = []
    for row in rows[4:]:
        if not row or not str(row[0]).startswith("BIA-"):
            continue
        rec = {h: clean(row[idx[h]]) if h in idx and idx[h] < len(row) else "" for h in idx}
        records.append(rec)
    if len(records) != 20:
        raise SystemExit(f"Expected 20 BIA rows, got {len(records)}")
    return records


def comments_for(rec):
    bits = ["No practice drill has been run yet."]
    if rec["Cross-Trained Backup Staff Identified?"] == "No":
        bits.append("Backup-staff gap must close before next review.")
    return " ".join(bits)


def strategy_rows(bia_records):
    bad = [v for v in STRATEGY_TYPE.values() if v not in ALLOWED_STRATEGY_TYPES]
    if bad:
        raise SystemExit(f"Strategy Type not in dropdown: {bad}")
    missing = [rec["Function / Process Name"] for rec in bia_records if rec["Function / Process Name"] not in STRATEGY_TYPE]
    if missing:
        raise SystemExit(f"No strategy type for: {missing}")
    out = []
    for i, rec in enumerate(bia_records, start=1):
        name = rec["Function / Process Name"]
        out.append(
            [
                f"RS-2026-TU-{i:02d}",
                rec["BIA Ref"],
                rec["Department / Unit"],
                name,
                rec["Completed By"] or "R. Sulaimon (BCC)",
                rec["Profiling Date"] or "01-Sep-2026",
                "v1.0",
                rec["Overall Criticality Ranking"],
                rec["RTO (Hrs)"],
                rec["RPO"],
                rec["MTPD (Hrs)"],
                STRATEGY_TYPE[name],
                rec["Recommended Recovery Strategy"],
                rec["Alternate Site/Access Required?"],
                rec["Cross-Trained Backup Staff Identified?"],
                rec["Current Recovery Capability (Hrs)"],
                rec["RTO Gap?"],
                "Yes",
                RESOURCES_DOCUMENTED[name],
                "",  # Last Tested Date — no drill yet
                "Not yet tested",
                "",  # Next Test Due — no date
                "Not Tested",
                "",
                "",
                "",
                "",
                "",
                "",
                comments_for(rec),
            ]
        )
    return out


def write_csv(data_rows):
    with TEMPLATE_PATH.open(newline="", encoding="latin-1") as f:
        template = list(csv.reader(f))
    # Keep title, department line, group headers, column headers. Replace sample rows.
    header_block = template[:4]
    # Fill department on row 2
    dept_row = header_block[1]
    if dept_row:
        dept_row[0] = "Department: Implementation / Technical Unit"
        # keep Register Owner placeholder on the same row if present
        for i, cell in enumerate(dept_row):
            if "Register Owner" in cell:
                dept_row[i] = (
                    "Register Owner (BCC):  ______________________     "
                    "Last Updated:  03-Sep-2026"
                )
    with OUT_CSV.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerows(header_block)
        writer.writerows(data_rows)
    print(f"Wrote {OUT_CSV} ({len(data_rows)} strategies)")


COLUMN_HEADERS = [
    "Strategy Ref",
    "BIA Ref",
    "Department / Unit",
    "Function / Process Name",
    "Prepared By (BCC)",
    "Date Prepared",
    "Plan Version",
    "Criticality Ranking",
    "RTO (Hrs)",
    "RPO",
    "MTPD (Hrs)",
    "Strategy Type(s)",
    "Strategy Summary",
    "Alternate Site/Access Required?",
    "Cross-Trained Backup Staff?",
    "Current Recovery Capability (Hrs)",
    "RTO Gap?",
    "Recovery Team Confirmed?",
    "Resource Requirements Documented?",
    "Last Tested Date",
    "Test Result Summary",
    "Next Test Due",
    "Test Status",
    "Reviewed By",
    "Review Date",
    "Approved By",
    "Approval Date",
    "Status",
    "Days Since Last Review",
    "Comments / Notes",
]

NAVY_FILL = PatternFill("solid", fgColor=NAVY)
TEAL_FILL = PatternFill("solid", fgColor=TEAL)
WHITE_FONT = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
TITLE_FONT = Font(name="Calibri", size=16, bold=True, color=NAVY)
BODY_FONT = Font(name="Calibri", size=10, color="1A1A1A")
THIN = Border(
    left=Side(style="thin", color="C5CDD6"),
    right=Side(style="thin", color="C5CDD6"),
    top=Side(style="thin", color="C5CDD6"),
    bottom=Side(style="thin", color="C5CDD6"),
)
WRAP = Alignment(wrap_text=True, vertical="center", horizontal="left")
CENTER = Alignment(wrap_text=True, vertical="center", horizontal="center")
GROUP_SPANS = [
    (1, 7, "IDENTIFICATION", NAVY_FILL),
    (8, 11, "RECOVERY TARGETS (FROM BIA)", TEAL_FILL),
    (12, 15, "STRATEGY DETAILS", NAVY_FILL),
    (16, 19, "READINESS", TEAL_FILL),
    (20, 23, "TESTING", NAVY_FILL),
    (24, 30, "SIGN-OFF & STATUS", NAVY_FILL),
]


def write_xlsx(data_rows):
    wb = Workbook()
    lists = wb.active
    lists.title = "Lookup Lists"
    lists["A1"] = "Strategy Type(s)"
    lists["A1"].font = Font(name="Calibri", size=11, bold=True, color=NAVY)
    types = [
        "People",
        "Process",
        "Technology",
        "Facilities",
        "People + Process",
        "People + Technology",
        "Technology + Facilities",
        "People + Process + Technology",
    ]
    for i, value in enumerate(types, start=2):
        lists[f"A{i}"] = value
        lists[f"A{i}"].font = BODY_FONT
        lists[f"A{i}"].border = THIN
    lists["C1"] = "Yes / No"
    lists["C1"].font = Font(name="Calibri", size=11, bold=True, color=NAVY)
    lists["C2"] = "Yes"
    lists["C3"] = "No"
    lists["E1"] = "Test Status"
    lists["E1"].font = Font(name="Calibri", size=11, bold=True, color=NAVY)
    for i, value in enumerate(("Not Tested", "On Track", "Overdue"), start=2):
        lists[f"E{i}"] = value
        lists[f"E{i}"].font = BODY_FONT
        lists[f"E{i}"].border = THIN
    lists.column_dimensions["A"].width = 36
    lists.column_dimensions["C"].width = 14
    lists.column_dimensions["E"].width = 16
    wb.defined_names.add(DefinedName(name="StrategyTypes", attr_text="'Lookup Lists'!$A$2:$A$9"))
    wb.defined_names.add(DefinedName(name="YesNo", attr_text="'Lookup Lists'!$C$2:$C$3"))
    wb.defined_names.add(DefinedName(name="TestStatus", attr_text="'Lookup Lists'!$E$2:$E$4"))

    ws = wb.create_sheet("Recovery Strategy Register", 0)
    ws.merge_cells("A1:AD1")
    ws["A1"] = "DEPARTMENTAL RECOVERY STRATEGY REGISTER"
    ws["A1"].font = TITLE_FONT
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 24

    ws.merge_cells("A2:G2")
    ws["A2"] = "Department: Implementation / Technical Unit"
    ws["A2"].font = Font(name="Calibri", size=11, bold=True, color=NAVY)
    ws.merge_cells("H2:AD2")
    ws["H2"] = "Register Owner (BCC):  ______________________     Last Updated:  03-Sep-2026"
    ws["H2"].font = Font(name="Calibri", size=10, italic=True, color="5A6A7A")
    ws.row_dimensions[2].height = 20

    for start, end, label, fill in GROUP_SPANS:
        ws.merge_cells(start_row=3, start_column=start, end_row=3, end_column=end)
        cell = ws.cell(3, start, label)
        cell.fill = fill
        cell.font = WHITE_FONT
        cell.alignment = CENTER
        for col in range(start, end + 1):
            ws.cell(3, col).fill = fill
            ws.cell(3, col).border = THIN
    ws.row_dimensions[3].height = 20

    for col, header in enumerate(COLUMN_HEADERS, start=1):
        cell = ws.cell(4, col, header)
        cell.font = WHITE_FONT
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
        cell.border = THIN
        cell.fill = TEAL_FILL if 8 <= col <= 11 or 16 <= col <= 19 else NAVY_FILL
    ws.row_dimensions[4].height = 36

    first = 5
    last = first + len(data_rows) - 1
    center_cols = {5, 6, 7, 8, 9, 10, 11, 12, 14, 15, 16, 17, 18, 19, 23}
    for r_i, row in enumerate(data_rows):
        r = first + r_i
        zebra = PatternFill("solid", fgColor="F4F7FA") if r_i % 2 else PatternFill("solid", fgColor="FFFFFF")
        for c_i, value in enumerate(row, start=1):
            cell = ws.cell(r, c_i, value)
            cell.font = Font(name="Calibri", size=10, bold=(c_i in (1, 4)), color="1A1A1A")
            cell.alignment = CENTER if c_i in center_cols else WRAP
            cell.border = THIN
            cell.fill = zebra
        ws.row_dimensions[r].height = 48

    dv_type = DataValidation(
        type="list",
        formula1="=StrategyTypes",
        allow_blank=False,
        showDropDown=False,
        showErrorMessage=True,
        errorTitle="Invalid strategy type",
        error="Choose a value from the Strategy Type(s) dropdown.",
    )
    dv_yes = DataValidation(type="list", formula1="=YesNo", allow_blank=True, showDropDown=False)
    dv_test = DataValidation(type="list", formula1="=TestStatus", allow_blank=True, showDropDown=False)
    ws.add_data_validation(dv_type)
    ws.add_data_validation(dv_yes)
    ws.add_data_validation(dv_test)
    dv_type.add(f"L{first}:L{last}")
    for col in ("N", "O", "Q", "R", "S"):
        dv_yes.add(f"{col}{first}:{col}{last}")
    dv_test.add(f"W{first}:W{last}")

    widths = {
        "A": 16, "B": 16, "C": 28, "D": 34, "E": 20, "F": 14, "G": 12,
        "H": 18, "I": 12, "J": 14, "K": 12, "L": 30, "M": 48, "N": 18,
        "O": 22, "P": 18, "Q": 12, "R": 18, "S": 22, "T": 16, "U": 22,
        "V": 14, "W": 14, "X": 16, "Y": 14, "Z": 16, "AA": 14, "AB": 14,
        "AC": 16, "AD": 42,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    ws.freeze_panes = "E5"
    ws.auto_filter.ref = f"A4:AD{last}"
    ws.sheet_properties.tabColor = NAVY
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 4
    ws.print_title_rows = "1:4"

    wb.save(OUT_XLSX)
    print(f"Wrote {OUT_XLSX}")


def set_run(run, *, size=10, bold=False, color=BODY, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "8AA0B4")
        b.append(el)
    tcPr.append(b)


def fill_cell(cell, text, *, bold=False, size=7, fill=None, font_color=BODY, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=font_color)
    borders(cell)
    if fill:
        shade(cell, fill)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    cell._tc.get_or_add_tcPr().append(v)


def write_docx(data_rows):
    # Readable extract: the strategy working columns, not all 30 Excel fields.
    headers = [
        "Strategy Ref",
        "BIA Ref",
        "Function / Process Name",
        "Criticality",
        "RTO",
        "RPO",
        "MTPD",
        "Strategy Type",
        "Strategy Summary",
        "Alternate Site/Access?",
        "Cross-Trained Backup?",
        "Current Recovery (Hrs)",
        "RTO Gap?",
        "Test Result Summary",
        "Test Status",
        "Comments / Notes",
    ]
    # indices in data_rows
    pick = [0, 1, 3, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 20, 22, 29]
    widths = [2.4, 2.8, 3.6, 2.4, 1.5, 2.2, 1.6, 2.8, 5.4, 2.4, 2.6, 2.4, 1.8, 2.6, 2.2, 3.6]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)

    title = doc.add_paragraph()
    r = title.add_run("Departmental Recovery Strategy Register")
    set_run(r, size=16, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    r = sub.add_run(
        "Department: Implementation / Technical Unit. One recovery strategy per BIA "
        "function. Recovery targets (criticality, RTO, RPO, MTPD) and the strategy "
        "summary are taken from the final Technical BIA. Test Result Summary is "
        "“Not yet tested” on every row — that is not a date and not a prediction "
        "that the risk will happen. Sign-off is left blank for the workshop. "
        "The full Excel/CSV has every template column."
    )
    set_run(r, size=10, color=BODY)

    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, label in enumerate(headers):
        fill_cell(hdr.cells[i], label, bold=True, size=7, fill=NAVY, font_color=WHITE, center=True)
        hdr.cells[i].width = Cm(widths[i])
    for r_i, row in enumerate(data_rows):
        zebra = "F4F7FA" if r_i % 2 else WHITE
        cells = table.rows[r_i + 1].cells
        for c_i, src in enumerate(pick):
            fill_cell(
                cells[c_i],
                row[src],
                bold=(c_i == 0 or c_i == 2),
                size=7,
                fill=zebra,
                center=c_i in (3, 4, 5, 6, 9, 10, 11, 12, 13, 14),
            )
            cells[c_i].width = Cm(widths[c_i])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    r = note.add_run(
        "Open Technical_Department_Recovery_Strategy_Register.xlsx in Excel to see "
        "the full template with the Strategy Type(s) dropdown."
    )
    set_run(r, size=9, color="5A6A7A")
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


def main():
    bia = load_bia()
    data = strategy_rows(bia)
    write_csv(data)
    write_xlsx(data)
    write_docx(data)


if __name__ == "__main__":
    main()
